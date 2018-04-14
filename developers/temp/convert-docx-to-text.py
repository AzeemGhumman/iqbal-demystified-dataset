from docx import Document
import sys
import glob
import os
import docx
import ntpath
import pdb

if (len(sys.argv) is not 3):
    print ("Error: Need <input folder> followed by <output folder>")
    exit()

inputFolder = sys.argv[1] # first argument: input folder
outputFolder = sys.argv[2] # second argument: output folder

# Get list of .docx files from the input folder
listInputFiles = glob.glob(os.path.join(os.getcwd(), inputFolder, "*.docx"))
print (str(len(listInputFiles)) + " input files found..")

for inputFilename in listInputFiles:

    wordDoc = Document(inputFilename)
    print ("Processing " + ntpath.basename(inputFilename))

    urduHeading = wordDoc.paragraphs[1].text
    englishHeading = wordDoc.paragraphs[3].text

    fileContents = []
    fileContents.append(urduHeading)
    fileContents.append("*")

    foundText = False
    for table in wordDoc.tables:
        for row in table.rows:
            if "Urdu Text" in row.cells[0].text:
                foundText = True
                cellContents = row.cells[1].text
                lines = cellContents.split("\n")
                if len(lines) != 2:
                    print ("Error: Invalid number of lines in sher")
                    pdb.set_trace()
                fileContents.extend(lines)
        if foundText is False:
            print ("Error: Could not find Urdu text in this table")

    fileContents.append("*")
    fileContents.append(englishHeading)
    fileContents.append("*")
    fileContents.append("")
    fileContents.append("*")
    fileContents.append("")
    fileContents.append("*")
    fileContents.append("")
    fileContents.append("*")
    fileContents.append("")
    fileContents.append("*")
    fileContents.append("")
    fileContents.append("*")

    fileName = ntpath.basename(inputFilename)[:-5]
    outputFilePath = os.path.join(outputFolder, fileName + ".txt")

    # Create output folder if it does not exist
    if not os.path.exists(os.path.dirname(outputFilePath)):
        os.makedirs(os.path.dirname(outputFilePath))

    with open(outputFilePath, 'w') as outputFile:
        outputFile.write("\n".join(fileContents))
